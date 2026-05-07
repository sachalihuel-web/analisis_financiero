from docx import Document

def exportar_word(activo, retorno, volatilidad, riesgo, ruta_grafico):
    doc = Document()

    doc.add_heading(f'Informe Financiero - {activo}', 0)

    doc.add_heading('Métricas', level=1)
    doc.add_paragraph(f'Retorno: {retorno:.2%}')
    doc.add_paragraph(f'Volatilidad: {volatilidad:.2%}')
    doc.add_paragraph(f'Riesgo: {riesgo}')

    doc.add_heading('Gráfico', level=1)
    doc.add_picture(ruta_grafico)

    doc.add_heading('Conclusión', level=1)
    doc.add_paragraph(
        f'El activo presenta un nivel de riesgo {riesgo} con un retorno de {retorno:.2%}.'
    )

    ruta = f"reports/informe_{activo}.docx"
    doc.save(ruta)

    return ruta